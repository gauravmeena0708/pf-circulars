"""Generates formatted Microsoft Word (.docx) documents for research reports,
interactive Q&A transcripts, and extracted OCR text.
"""

from __future__ import annotations

from datetime import datetime
import io
import re
from typing import Any, Mapping, Sequence

import docx
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor


NAVY_COLOR = RGBColor(26, 54, 93)  # #1A365D
GRAY_COLOR = RGBColor(100, 116, 139)  # #64748B
DARK_TEXT_COLOR = RGBColor(30, 41, 59)  # #1E293B


def _set_cell_background(cell, fill_hex: str):
    """Sets the background color of a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tc_pr.append(shd)


def _set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets inner padding/margins for a table cell in dxa (1/20 pt)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tc_pr.append(tc_mar)


def _add_formatted_text(paragraph, text: str):
    """Parses simple inline markdown (bold **text**, italics *text*) into Word runs."""
    # Split by bold markers **...**
    tokens = re.split(r"(\*\*.*?\*\*)", text)
    for token in tokens:
        if token.startswith("**") and token.endswith("**") and len(token) >= 4:
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            # Check for *italic*
            sub_tokens = re.split(r"(\*.*?\*)", token)
            for st in sub_tokens:
                if st.startswith("*") and st.endswith("*") and len(st) >= 2:
                    run = paragraph.add_run(st[1:-1])
                    run.italic = True
                else:
                    paragraph.add_run(st)


def _render_markdown_blocks(doc: docx.Document, content: str):
    """Appends paragraphs, bullet points, and headers from a markdown string."""
    lines = content.split("\n")
    in_code_block = False

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            continue

        if not stripped:
            continue

        if in_code_block:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = GRAY_COLOR
            continue

        if stripped.startswith("### "):
            p = doc.add_heading(level=3)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            _add_formatted_text(p, stripped[4:])
        elif stripped.startswith("## "):
            p = doc.add_heading(level=2)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            _add_formatted_text(p, stripped[3:])
        elif stripped.startswith("# "):
            p = doc.add_heading(level=1)
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(6)
            _add_formatted_text(p, stripped[2:])
        elif stripped.startswith(("- ", "* ", "• ")):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            _add_formatted_text(p, stripped[2:])
        elif re.match(r"^\d+\.\s", stripped):
            match = re.match(r"^\d+\.\s", stripped)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(3)
            _add_formatted_text(p, stripped[match.end():])
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            _add_formatted_text(p, stripped)


def create_research_report_docx(
    query: str,
    answer_text: str,
    source_references: Sequence[Mapping[str, Any]],
    generated_at: str | None = None,
) -> bytes:
    """
    Creates a styled Word (.docx) document for an EPFO research & citation report.
    """
    doc = docx.Document()

    # Configure Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Document Header Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("EPFO Knowledge Assistant — Research & Citation Report")
    run_title.font.name = "Calibri"
    run_title.font.size = Pt(20)
    run_title.bold = True
    run_title.font.color.rgb = NAVY_COLOR

    # Timestamp & Metadata Subtitle
    date_str = generated_at or datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(12)
    sub_run = sub_p.add_run(f"Generated on {date_str}")
    sub_run.font.size = Pt(9.5)
    sub_run.font.color.rgb = GRAY_COLOR

    # Query Card / Box
    q_table = doc.add_table(rows=1, cols=1)
    q_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    q_table.autofit = False
    cell = q_table.rows[0].cells[0]
    cell.width = Inches(6.5)
    _set_cell_background(cell, "F1F5F9")  # Slate-100
    _set_cell_margins(cell, top=140, bottom=140, left=180, right=180)

    qp = cell.paragraphs[0]
    qp.paragraph_format.space_after = Pt(0)
    q_label = qp.add_run("Research Query: ")
    q_label.bold = True
    q_label.font.color.rgb = NAVY_COLOR
    qp.add_run(query)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Section 1: Synthesized Answer
    h1 = doc.add_heading(level=1)
    h1_run = h1.add_run("💡 Synthesized Answer & Analysis")
    h1_run.font.color.rgb = NAVY_COLOR
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)

    if answer_text and answer_text.strip():
        _render_markdown_blocks(doc, answer_text.strip())
    else:
        empty_p = doc.add_paragraph()
        empty_run = empty_p.add_run("No AI synthesized answer generated (Search & Citations mode).")
        empty_run.italic = True
        empty_run.font.color.rgb = GRAY_COLOR

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Section 2: Source References & Citations
    if source_references:
        h2 = doc.add_heading(level=1)
        h2_run = h2.add_run(f"📚 Source References ({len(source_references)})")
        h2_run.font.color.rgb = NAVY_COLOR
        h2.paragraph_format.space_before = Pt(14)
        h2.paragraph_format.space_after = Pt(8)

        for i, ref in enumerate(source_references, start=1):
            ref_box = doc.add_table(rows=1, cols=1)
            ref_box.alignment = WD_TABLE_ALIGNMENT.CENTER
            ref_cell = ref_box.rows[0].cells[0]
            ref_cell.width = Inches(6.5)
            _set_cell_background(ref_cell, "F8FAFC")
            _set_cell_margins(ref_cell, top=120, bottom=120, left=160, right=160)

            rp = ref_cell.paragraphs[0]
            rp.paragraph_format.space_after = Pt(3)

            title = ref.get("title") or ref.get("subject") or f"Reference #{i}"
            source = ref.get("source") or ref.get("circular_no") or "Official EPFO Document"
            date = ref.get("date") or ref.get("circular_date") or "N/A"
            score = ref.get("score")

            title_run = rp.add_run(f"#{i} {title}\n")
            title_run.bold = True
            title_run.font.size = Pt(11)
            title_run.font.color.rgb = NAVY_COLOR

            meta_line = f"Source: {source} | Date: {date}"
            if score is not None:
                try:
                    meta_line += f" | Relevance: {float(score):.2f}"
                except (ValueError, TypeError):
                    pass

            meta_p = ref_cell.add_paragraph()
            meta_p.paragraph_format.space_after = Pt(4)
            m_run = meta_p.add_run(meta_line)
            m_run.font.size = Pt(9)
            m_run.font.color.rgb = GRAY_COLOR

            snippet = ref.get("text") or ref.get("snippet") or ""
            if snippet:
                snip_p = ref_cell.add_paragraph()
                snip_p.paragraph_format.space_after = Pt(2)
                s_run = snip_p.add_run(f'"{snippet.strip()}"')
                s_run.italic = True
                s_run.font.size = Pt(9.5)

            url = ref.get("url") or ref.get("link")
            if url:
                url_p = ref_cell.add_paragraph()
                url_p.paragraph_format.space_after = Pt(0)
                u_run = url_p.add_run(f"Official Link: {url}")
                u_run.font.size = Pt(8.5)
                u_run.font.color.rgb = RGBColor(37, 99, 235)  # Blue

            doc.add_paragraph().paragraph_format.space_after = Pt(4)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def create_chat_transcript_docx(
    title: str,
    chat_history: Sequence[Mapping[str, str]],
    metadata: Mapping[str, Any] | None = None,
) -> bytes:
    """
    Creates a styled Word (.docx) document for conversational Q&A or data analysis sessions.
    """
    doc = docx.Document()

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Document Header Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run(title)
    run_title.font.name = "Calibri"
    run_title.font.size = Pt(18)
    run_title.bold = True
    run_title.font.color.rgb = NAVY_COLOR

    date_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(10)
    sub_run = sub_p.add_run(f"Generated on {date_str}")
    sub_run.font.size = Pt(9.5)
    sub_run.font.color.rgb = GRAY_COLOR

    # Optional Metadata Table
    if metadata:
        meta_table = doc.add_table(rows=len(metadata), cols=2)
        meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row_idx, (k, v) in enumerate(metadata.items()):
            row = meta_table.rows[row_idx]
            k_cell, v_cell = row.cells[0], row.cells[1]
            k_cell.width = Inches(2.0)
            v_cell.width = Inches(4.5)
            _set_cell_background(k_cell, "F1F5F9")
            _set_cell_background(v_cell, "F8FAFC")
            _set_cell_margins(k_cell, 80, 80, 100, 100)
            _set_cell_margins(v_cell, 80, 80, 100, 100)

            kp = k_cell.paragraphs[0]
            k_run = kp.add_run(str(k))
            k_run.bold = True
            k_run.font.size = Pt(9.5)

            vp = v_cell.paragraphs[0]
            v_run = vp.add_run(str(v))
            v_run.font.size = Pt(9.5)

        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Render Chat Conversation
    for i, msg in enumerate(chat_history):
        role = msg.get("role", "user").lower()
        content = msg.get("content", "")

        is_user = role == "user"
        role_label = "👤 User Query" if is_user else "🤖 Analysis & Response"

        head = doc.add_heading(level=2)
        head.paragraph_format.space_before = Pt(12)
        head.paragraph_format.space_after = Pt(4)
        h_run = head.add_run(role_label)
        h_run.font.color.rgb = RGBColor(30, 64, 175) if is_user else NAVY_COLOR

        _render_markdown_blocks(doc, content)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def create_text_document_docx(
    title: str,
    body_text: str,
    metadata: Mapping[str, Any] | None = None,
) -> bytes:
    """
    Creates a styled Word (.docx) document from raw extracted text.
    """
    doc = docx.Document()

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run(title)
    run_title.font.name = "Calibri"
    run_title.font.size = Pt(18)
    run_title.bold = True
    run_title.font.color.rgb = NAVY_COLOR

    date_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(12)
    sub_run = sub_p.add_run(f"Exported on {date_str}")
    sub_run.font.size = Pt(9.5)
    sub_run.font.color.rgb = GRAY_COLOR

    if metadata:
        for k, v in metadata.items():
            mp = doc.add_paragraph()
            mp.paragraph_format.space_after = Pt(2)
            k_run = mp.add_run(f"{k}: ")
            k_run.bold = True
            mp.add_run(str(v))
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    _render_markdown_blocks(doc, body_text)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
