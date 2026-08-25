import io
import unittest
import docx

import docx_export


class TestDocxExport(unittest.TestCase):

    def test_create_research_report_docx(self):
        query = "What is the procedure for joint declaration?"
        answer = "To update joint declaration, the member and employer must submit the request online via Unified Portal."
        references = [
            {
                "rank": 1,
                "title": "Standard Operating Procedure for Joint Declaration",
                "source": "Circular No. 12345",
                "date": "2023-10-15",
                "score": 0.92,
                "text": "The joint declaration application should be verified by the employer within 30 days.",
                "url": "https://epfindia.gov.in/circulars/12345.pdf",
            }
        ]

        doc_bytes = docx_export.create_research_report_docx(query, answer, references)
        self.assertIsInstance(doc_bytes, bytes)
        self.assertGreater(len(doc_bytes), 1000)

        # Verify it can be loaded by python-docx
        doc = docx.Document(io.BytesIO(doc_bytes))
        all_text = "\n".join(p.text for p in doc.paragraphs) + "\n" + "\n".join(
            c.text for t in doc.tables for r in t.rows for c in r.cells
        )
        self.assertIn("EPFO Knowledge Assistant", all_text)
        self.assertIn(query, all_text)
        self.assertIn("Standard Operating Procedure", all_text)

    def test_create_chat_transcript_docx(self):
        title = "Data Analysis: employee_data.csv"
        chat_history = [
            {"role": "user", "content": "What is the average salary in the dataset?"},
            {"role": "assistant", "content": "The average salary observed is ₹45,000 based on 1,200 records."},
        ]
        metadata = {"File Name": "employee_data.csv", "Total Rows": 1200}

        doc_bytes = docx_export.create_chat_transcript_docx(title, chat_history, metadata)
        self.assertIsInstance(doc_bytes, bytes)
        self.assertGreater(len(doc_bytes), 1000)

        doc = docx.Document(io.BytesIO(doc_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn(title, full_text)
        self.assertIn("What is the average salary", full_text)
        self.assertIn("45,000", full_text)

    def test_create_text_document_docx(self):
        title = "Extracted OCR Text"
        body = "--- [Page 1 / 1] ---\nThis is scanned noting sheet text from regional office."
        doc_bytes = docx_export.create_text_document_docx(title, body)
        self.assertIsInstance(doc_bytes, bytes)
        self.assertGreater(len(doc_bytes), 500)

        doc = docx.Document(io.BytesIO(doc_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn(title, full_text)
        self.assertIn("scanned noting sheet text", full_text)


if __name__ == "__main__":
    unittest.main()
